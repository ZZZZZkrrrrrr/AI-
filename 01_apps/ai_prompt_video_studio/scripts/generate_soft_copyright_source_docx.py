from __future__ import annotations

import unicodedata
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "copyright_materials"
OUTPUT_FILE = OUTPUT_DIR / "software_copyright_source_identification_V1.0_20260613_stable.docx"

SOFTWARE_NAME = "AI短视频智能创作工作流管理系统软件"
VERSION = "V1.0"
LINES_PER_PAGE = 50
PAGES_PER_PART = 30
TOTAL_PAGES = PAGES_PER_PART * 2
MAX_DISPLAY_WIDTH = 220


def source_files() -> list[Path]:
    preferred = [
        "src/App.jsx",
        "src/main.jsx",
        "src/api.js",
        "src/features/auth/LoginPromptSheet.jsx",
        "src/features/assets/AssetsPage.jsx",
        "src/features/batch/BatchPage.jsx",
        "src/features/settings/SettingsPage.jsx",
        "src/features/stitch/VideoStitchPage.jsx",
        "src/features/textImage/TextImagePage.jsx",
        "src/features/workflow/WorkflowModulePage.jsx",
        "src/shared/compliance/aiEvidencePack.js",
        "src/shared/pwa/registerServiceWorker.js",
        "src/shared/textImageStudioHandoff.js",
        "src/styles.css",
        "server.js",
    ]
    return [ROOT / item for item in preferred if (ROOT / item).exists()]


def read_source_lines(files: list[Path]) -> list[str]:
    lines: list[str] = []
    for path in files:
        relative = path.relative_to(ROOT).as_posix()
        lines.append(f"/* ===== File: {relative} ===== */")
        text = path.read_text(encoding="utf-8", errors="replace")
        for line in text.splitlines():
            lines.append(line.expandtabs(2).rstrip())
        lines.append(f"/* ===== End File: {relative} ===== */")
        lines.append("")
    return lines


def selected_lines(all_lines: list[str]) -> list[tuple[str, int, str]]:
    front_count = PAGES_PER_PART * LINES_PER_PAGE
    back_count = PAGES_PER_PART * LINES_PER_PAGE
    if len(all_lines) <= front_count + back_count:
        selected = [("源程序全文", index, line) for index, line in enumerate(all_lines, start=1)]
    else:
        selected = [
            ("源程序前连续30页", index, line)
            for index, line in enumerate(all_lines[:front_count], start=1)
        ]
        back_start = len(all_lines) - back_count + 1
        selected.extend(
            ("源程序后连续30页", index, line)
            for index, line in enumerate(all_lines[-back_count:], start=back_start)
        )
    target = TOTAL_PAGES * LINES_PER_PAGE
    if len(selected) < target:
        selected.extend(("补空", 0, "") for _ in range(target - len(selected)))
    return selected[:target]


def display_width(text: str) -> int:
    width = 0
    for char in text:
        if unicodedata.east_asian_width(char) in {"F", "W", "A"}:
            width += 2
        else:
            width += 1
    return width


def clamp_display(text: str, max_width: int) -> str:
    width = 0
    out: list[str] = []
    for char in text:
        char_width = 2 if unicodedata.east_asian_width(char) in {"F", "W", "A"} else 1
        if width + char_width > max_width:
            return "".join(out).rstrip() + " ..."
        out.append(char)
        width += char_width
    return "".join(out)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "C8D4CE") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def style_paragraph(paragraph, font_name: str = "SimSun", font_size: float = 5.6, color: str = "111111") -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(7.0)
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text: str, font_name: str = "SimSun", font_size: float = 5.6, bold: bool = False, color: str = "111111") -> None:
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def create_docx(rows: list[tuple[str, int, str]]) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.36)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.36)
    section.right_margin = Inches(0.36)
    section.header_distance = Inches(0.12)
    section.footer_distance = Inches(0.12)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(5.6)

    for page_index in range(TOTAL_PAGES):
        start = page_index * LINES_PER_PAGE
        page_rows = rows[start : start + LINES_PER_PAGE]
        page_no = page_index + 1
        part = page_rows[0][0] if page_rows else ""
        part_page = page_no if page_no <= PAGES_PER_PART else page_no - PAGES_PER_PART

        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(0)
        header.paragraph_format.space_after = Pt(1)
        header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        header.paragraph_format.line_spacing = Pt(9)
        add_text(header, f"{SOFTWARE_NAME} {VERSION} 源代码", font_size=7.4, bold=True, color="16211D")
        add_text(header, " " * 86, font_size=7.4)
        add_text(header, f"{part} 第{part_page:02d}页", font_size=7.4, color="16211D")

        table = doc.add_table(rows=LINES_PER_PAGE, cols=1)
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(10.86)

        for row_index, (_, source_line_number, text) in enumerate(page_rows):
            cell = table.rows[row_index].cells[0]
            cell.width = Inches(10.86)
            cell.margin_left = 0
            cell.margin_right = 0
            set_cell_border(cell, "FFFFFF")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(7.0)
            display = f"{source_line_number:04d}  {clamp_display(text, MAX_DISPLAY_WIDTH)}"
            add_text(paragraph, display, font_name="SimSun", font_size=5.55, color="111111")

        footer = doc.add_paragraph()
        footer.paragraph_format.space_before = Pt(2)
        footer.paragraph_format.space_after = Pt(0)
        footer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        footer.paragraph_format.line_spacing = Pt(8)
        add_text(footer, "材料类型：程序鉴别材料（一般交存）", font_size=6.8, color="52645C")
        add_text(footer, " " * 76, font_size=6.8)
        add_text(footer, f"第 {page_no} 页 / 共 {TOTAL_PAGES} 页", font_size=6.8, color="52645C")
        add_text(footer, " " * 80, font_size=6.8)
        add_text(footer, f"生成日期：{date.today().isoformat()}", font_size=6.8, color="52645C")

        if page_no != TOTAL_PAGES:
            footer.add_run().add_break(WD_BREAK.PAGE)

    doc.core_properties.title = f"{SOFTWARE_NAME} {VERSION} 程序鉴别材料"
    doc.core_properties.author = "软件著作权登记材料"
    doc.save(OUTPUT_FILE)


def main() -> None:
    files = source_files()
    lines = read_source_lines(files)
    rows = selected_lines(lines)
    create_docx(rows)
    print(OUTPUT_FILE)
    print(f"source_files={len(files)} source_lines={len(lines)} pages={TOTAL_PAGES}")


if __name__ == "__main__":
    main()
